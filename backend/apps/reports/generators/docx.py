"""
NABL-style Calibration Certificate DOCX Generator.

Generates a multi-page DOCX report matching the format used by
Legal Metrology Laboratories under the Government of India,
following OIML R 76-1:2006 test procedures for non-automatic
weighing instruments.
"""

import logging
from collections import defaultdict
from decimal import Decimal
from pathlib import Path
from typing import Any

from django.conf import settings

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

from apps.engine.constants import ComplianceStatus, TestType, EccentricityPosition
from apps.reports.generators import _build_report_context

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
PASS_COLOR = RGBColor(0x00, 0x64, 0x00)       # dark green
FAIL_COLOR = RGBColor(0x8B, 0x00, 0x00)       # dark red
GRAY_COLOR = RGBColor(0x66, 0x66, 0x66)
HEADER_BG = 'D9E2F3'                           # light blue header row
ALT_ROW_BG = 'FAFAFA'                          # alternating row shading
BORDER_COLOR = '000000'

# Page geometry (A4)
PAGE_WIDTH_CM = 21.0
LEFT_MARGIN_CM = 1.5
RIGHT_MARGIN_CM = 1.5
USABLE_WIDTH_CM = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM  # 18 cm

# ---------------------------------------------------------------------------
# Test-type display titles
# ---------------------------------------------------------------------------
TEST_TYPE_TITLES: dict[str, str] = {
    TestType.WEIGHING_PERFORMANCE: 'Weighing Performance Test',
    TestType.ECCENTRICITY: 'Eccentricity Test',
    TestType.REPEATABILITY: 'Repeatability Test',
    TestType.DISCRIMINATION: 'Discrimination Test',
    TestType.SENSITIVITY: 'Sensitivity Test',
    TestType.TARE: 'Tare Device Test',
    TestType.CREEP: 'Creep / Time Dependence Test',
    TestType.ZERO_RETURN: 'Zero Return Test',
    TestType.TEMPERATURE: 'Temperature Effect Test',
    TestType.TILT: 'Tilt Test',
    TestType.POWER_SUPPLY: 'Power Supply Variation Test',
    TestType.DURABILITY: 'Durability Test',
    TestType.SPAN_STABILITY: 'Span Stability Test',
    TestType.ZERO_TRACKING: 'Zero Tracking Test',
}

IMPORTANT_REMARKS = [
    'This Test Report shall not be reproduced except in full, without the written approval of the laboratory.',
    'The results reported relate only to the items tested / calibrated.',
    'The test results are valid at the time of testing under the stated environmental conditions.',
    'The expanded measurement uncertainty is stated as the standard measurement uncertainty multiplied '
    'by the coverage factor k = 2, corresponding to a coverage probability of approximately 95%.',
    'This report does not constitute a certificate of verification or approval of the instrument for '
    'trade use under the Legal Metrology Act, 2009 and rules made thereunder.',
    'Any alteration or modification to the instrument after the date of test shall invalidate this report.',
    'The instrument shall be re-verified at regular intervals as prescribed under the relevant rules.',
]

DOC_CONTROL_NO = 'LM-FMT-NAWI-01'
DOC_ISSUE_NO = '01'
DOC_ISSUE_DATE = '01.01.2026'
DOC_REV_NO = '00'


# ===================================================================
# Helper utilities
# ===================================================================

def _format_decimal(val: Decimal | None) -> str:
    """Format a Decimal for display, stripping trailing zeros."""
    if val is None:
        return '—'  # em dash
    normalized = val.normalize()
    if normalized == normalized.to_integral_value():
        return str(int(normalized))
    return str(normalized)


def _status_label(s: str) -> str:
    if s == ComplianceStatus.PASS:
        return 'PASS'
    if s == ComplianceStatus.FAIL:
        return 'FAIL'
    return 'N/A'


def _format_date_display(date_str: str) -> str:
    """Convert YYYY-MM-DD to DD.MM.YYYY for NABL display."""
    if not date_str:
        return '—'
    parts = date_str.split('-')
    if len(parts) == 3:
        return f'{parts[2]}.{parts[1]}.{parts[0]}'
    return date_str


# -------------------------------------------------------------------
# Cell / table formatting helpers
# -------------------------------------------------------------------

def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                      color=BORDER_COLOR, size='4'):
    """Apply single-line borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, enabled in [('top', top), ('bottom', bottom),
                          ('left', left), ('right', right)]:
        if enabled:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), size)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_table_borders(table, color=BORDER_COLOR, size='4'):
    """Set borders on every cell in *table*."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color=color, size=size)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Apply a background fill colour to *cell*."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_vertical_alignment(cell, align='center'):
    """Set vertical alignment of cell content."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_cell_width(cell, cm_val: float):
    """Set a fixed width on *cell* in centimetres."""
    width_emu = int(cm_val * 360000)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_emu))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _set_cell_margins(cell, top=0, bottom=0, left=57, right=57):
    """Set inner margins for *cell* (values in twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom),
                      ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _style_header_row(row, bg_color=HEADER_BG) -> None:
    """Style a table header row: bold, background colour, centred."""
    for cell in row.cells:
        _set_cell_shading(cell, bg_color)
        _set_cell_vertical_alignment(cell, 'center')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'


def _add_kv_row(table, label: str, value: str, label_width_cm=6.0) -> None:
    """Add a label : value row to a two-column table."""
    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]
    _set_cell_width(cell_label, label_width_cm)
    _set_cell_width(cell_value, USABLE_WIDTH_CM - label_width_cm)
    _set_cell_shading(cell_label, 'F2F2F2')

    p_label = cell_label.paragraphs[0]
    run_label = p_label.add_run(label)
    run_label.font.size = Pt(10)
    run_label.font.bold = True
    run_label.font.name = 'Calibri'

    p_value = cell_value.paragraphs[0]
    run_value = p_value.add_run(str(value) if value else '—')
    run_value.font.size = Pt(10)
    run_value.font.name = 'Calibri'


def _write_cell(cell, text: str, size=9, bold=False, font_name='Calibri',
                color: RGBColor | None = None,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                italic=False, mono=False):
    """Write styled text into a table cell, clearing any existing content."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Courier New' if mono else font_name
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return run


def _write_status_cell(cell, status_str: str):
    """Write a PASS / FAIL / N/A status cell with appropriate colour."""
    label = _status_label(status_str)
    if status_str == ComplianceStatus.PASS:
        _write_cell(cell, label, bold=True, color=PASS_COLOR)
    elif status_str == ComplianceStatus.FAIL:
        _write_cell(cell, label, bold=True, color=FAIL_COLOR)
    else:
        _write_cell(cell, label, italic=True, color=GRAY_COLOR)


def _apply_alternating_shading(table, start_row=1):
    """Apply alternating white / light-gray shading to data rows."""
    for idx, row in enumerate(table.rows[start_row:], start=0):
        if idx % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, ALT_ROW_BG)


def _add_horizontal_rule(doc):
    """Insert a horizontal line paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_centered_text(doc, text: str, size: int = 10, bold: bool = False,
                        color: RGBColor | None = None,
                        font_name: str = 'Calibri',
                        space_after: int | None = None,
                        space_before: int | None = None,
                        uppercase: bool = False):
    """Add a centred paragraph with the given styling."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    display = text.upper() if uppercase else text
    run = p.add_run(display)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return p


def _add_header_footer(doc, context: dict):
    """Configure the section header / footer with page numbers."""
    for section in doc.sections:
        # --- Header ---
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(
            f"Test Report No: {context['report_number']}  |  "
            f"{context['lab_name']}"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = GRAY_COLOR
        run.font.name = 'Calibri'

        # --- Footer ---
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document control line
        run = fp.add_run(
            f"Doc No: {context.get('doc_control_number', DOC_CONTROL_NO)}  |  "
            f"Issue No: {context.get('doc_issue_number', DOC_ISSUE_NO)}  |  "
            f"Issue Date: {DOC_ISSUE_DATE}  |  "
            f"Rev No: {context.get('doc_rev_number', DOC_REV_NO)}"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = GRAY_COLOR
        run.font.name = 'Calibri'

        # Page number (field code)
        fp2 = footer.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = fp2.add_run('Page ')
        run2.font.size = Pt(7)
        run2.font.color.rgb = GRAY_COLOR
        run2.font.name = 'Calibri'

        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_page = fp2.add_run()
        run_page._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run_page2 = fp2.add_run()
        run_page2._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_page3 = fp2.add_run()
        run_page3._r.append(fldChar2)

        run_of = fp2.add_run(' of ')
        run_of.font.size = Pt(7)
        run_of.font.color.rgb = GRAY_COLOR
        run_of.font.name = 'Calibri'

        # NUMPAGES field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run_np = fp2.add_run()
        run_np._r.append(fldChar3)
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = ' NUMPAGES '
        run_np2 = fp2.add_run()
        run_np2._r.append(instrText2)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run_np3 = fp2.add_run()
        run_np3._r.append(fldChar4)


# ===================================================================
# Per-test-type table builders
# ===================================================================

def _build_weighing_performance(doc, results, unit, section_label):
    """Weighing Performance: separate increasing / decreasing sub-tables."""
    increasing = [r for r in results
                  if r.observation and r.observation.direction == 'increasing']
    decreasing = [r for r in results
                  if r.observation and r.observation.direction == 'decreasing']
    no_dir = [r for r in results
              if not r.observation or r.observation.direction not in ('increasing', 'decreasing')]

    for direction_label, subset in [
        ('Increasing Load', increasing or no_dir),
        ('Decreasing Load', decreasing),
    ]:
        if not subset:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f'{direction_label}:')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = True
        run.font.name = 'Calibri'

        headers = ['Sr No', f'Test Point ({unit})', f'Reference Wt ({unit})',
                   f'Indicated ({unit})', f'Error ({unit})',
                   f'MPE (±{unit})', f'U ±{unit} (k=2)', 'Result']
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            _write_cell(hdr.cells[i], h, size=8, bold=True)
        _style_header_row(hdr)

        for idx, r in enumerate(subset, start=1):
            row = table.add_row()
            obs = r.observation
            indicated = _format_decimal(obs.indicated_value if obs else None)
            ref_val = _format_decimal(obs.reference_value if obs else None)
            mpe_str = f'±{_format_decimal(r.mpe_applicable)}' if r.mpe_applicable else '—'
            uncertainty = getattr(r, 'expanded_uncertainty', None)
            u_str = f'±{_format_decimal(uncertainty)}' if uncertainty else '—'

            _write_cell(row.cells[0], str(idx), mono=True)
            _write_cell(row.cells[1], _format_decimal(r.test_point_load), mono=True)
            _write_cell(row.cells[2], ref_val, mono=True)
            _write_cell(row.cells[3], indicated, mono=True)
            _write_cell(row.cells[4], _format_decimal(r.computed_error), mono=True)
            _write_cell(row.cells[5], mpe_str, mono=True)
            _write_cell(row.cells[6], u_str, mono=True)
            _write_status_cell(row.cells[7], r.compliance_status)

        _apply_alternating_shading(table)
        _set_table_borders(table)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _build_eccentricity(doc, results, unit, section_label):
    """Eccentricity: position-based table with max diff and overall result."""
    headers = ['Sr No', 'Position', f'Test Load ({unit})', f'Indicated ({unit})',
               f'Error ({unit})', f'MPE (±{unit})', 'Result']
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        _write_cell(hdr.cells[i], h, size=8, bold=True)
    _style_header_row(hdr)

    # Determine centre reading for max-diff calculation
    centre_error = None
    for r in results:
        pos = (r.position or '').lower()
        if pos == 'center' or pos == 'centre':
            centre_error = r.computed_error
            break

    max_diff = Decimal('0')
    for idx, r in enumerate(results, start=1):
        row = table.add_row()
        obs = r.observation
        indicated = _format_decimal(obs.indicated_value if obs else None)
        pos_display = (r.position or '').replace('_', ' ').title()
        mpe_str = f'±{_format_decimal(r.mpe_applicable)}' if r.mpe_applicable else '—'

        _write_cell(row.cells[0], str(idx), mono=True)
        _write_cell(row.cells[1], pos_display)
        _write_cell(row.cells[2], _format_decimal(r.test_point_load), mono=True)
        _write_cell(row.cells[3], indicated, mono=True)
        _write_cell(row.cells[4], _format_decimal(r.computed_error), mono=True)
        _write_cell(row.cells[5], mpe_str, mono=True)
        _write_status_cell(row.cells[6], r.compliance_status)

        if centre_error is not None and r.computed_error is not None:
            diff = abs(r.computed_error - centre_error)
            if diff > max_diff:
                max_diff = diff

    _apply_alternating_shading(table)
    _set_table_borders(table)

    # Max difference summary row
    if centre_error is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f'Maximum difference from centre reading: {_format_decimal(max_diff)} {unit}')
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _build_repeatability(doc, results, unit, section_label):
    """Repeatability: grouped by test load, trials shown horizontally."""
    # Group by test_point_load
    by_load: dict[Decimal, list] = defaultdict(list)
    for r in results:
        load_key = r.test_point_load if r.test_point_load is not None else Decimal('0')
        by_load[load_key].append(r)

    for load_val in sorted(by_load.keys()):
        trials = sorted(by_load[load_val], key=lambda r: r.trial_number or 0)
        n_trials = len(trials)
        if n_trials == 0:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f'Test Load: {_format_decimal(load_val)} {unit}')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = True
        run.font.name = 'Calibri'

        # Header row: Trial labels
        trial_headers = [f'T{t.trial_number}' for t in trials]
        headers = ['Parameter'] + trial_headers + ['Range', f'MPE (±{unit})', 'Result']
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            _write_cell(hdr.cells[i], h, size=8, bold=True)
        _style_header_row(hdr)

        # Indicated values row
        row_ind = table.add_row()
        _write_cell(row_ind.cells[0], f'Indicated ({unit})', size=8, bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
        indicated_values = []
        for j, t in enumerate(trials):
            obs = t.observation
            val = obs.indicated_value if obs else None
            _write_cell(row_ind.cells[1 + j], _format_decimal(val), mono=True)
            if val is not None:
                indicated_values.append(val)

        # Compute range
        if indicated_values:
            range_val = max(indicated_values) - min(indicated_values)
        else:
            range_val = None

        mpe_val = trials[0].mpe_applicable if trials else None
        mpe_str = f'±{_format_decimal(mpe_val)}' if mpe_val else '—'

        # Worst compliance across trials
        worst = ComplianceStatus.PASS
        for t in trials:
            if t.compliance_status == ComplianceStatus.FAIL:
                worst = ComplianceStatus.FAIL
                break

        _write_cell(row_ind.cells[1 + n_trials], _format_decimal(range_val), mono=True)
        _write_cell(row_ind.cells[2 + n_trials], mpe_str, mono=True)
        _write_status_cell(row_ind.cells[3 + n_trials], worst)

        # Error values row
        row_err = table.add_row()
        _write_cell(row_err.cells[0], f'Error ({unit})', size=8, bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for j, t in enumerate(trials):
            _write_cell(row_err.cells[1 + j], _format_decimal(t.computed_error), mono=True)
        # Leave range / MPE / result cells empty for the error row
        for k in range(n_trials + 1, len(headers)):
            _write_cell(row_err.cells[k], '')

        _apply_alternating_shading(table)
        _set_table_borders(table)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _build_discrimination(doc, results, unit, section_label):
    """Discrimination test table."""
    headers = ['Sr No', f'Test Load ({unit})', f'Indication Before I ({unit})',
               f'Indication After I′ ({unit})', f'Change I′−I ({unit})',
               'Result']
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        _write_cell(hdr.cells[i], h, size=8, bold=True)
    _style_header_row(hdr)

    for idx, r in enumerate(results, start=1):
        row = table.add_row()
        obs = r.observation
        indicated = _format_decimal(obs.indicated_value if obs else None)
        ref_val = _format_decimal(obs.reference_value if obs else None)
        # Change = indicated - reference (before -> after adding 1.4d)
        change = '—'
        if obs and obs.indicated_value is not None and obs.reference_value is not None:
            change = _format_decimal(obs.indicated_value - obs.reference_value)

        _write_cell(row.cells[0], str(idx), mono=True)
        _write_cell(row.cells[1], _format_decimal(r.test_point_load), mono=True)
        _write_cell(row.cells[2], ref_val, mono=True)
        _write_cell(row.cells[3], indicated, mono=True)
        _write_cell(row.cells[4], change, mono=True)
        _write_status_cell(row.cells[5], r.compliance_status)

    _apply_alternating_shading(table)
    _set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _build_generic_test(doc, results, unit, section_label):
    """Generic test table for all other test types."""
    headers = ['Sr No', f'Test Point ({unit})', f'Indicated ({unit})',
               f'Error ({unit})', f'MPE (±{unit})', 'Result']
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        _write_cell(hdr.cells[i], h, size=8, bold=True)
    _style_header_row(hdr)

    for idx, r in enumerate(results, start=1):
        row = table.add_row()
        obs = r.observation
        indicated = _format_decimal(obs.indicated_value if obs else None)
        mpe_str = f'±{_format_decimal(r.mpe_applicable)}' if r.mpe_applicable else '—'

        _write_cell(row.cells[0], str(idx), mono=True)
        _write_cell(row.cells[1], _format_decimal(r.test_point_load), mono=True)
        _write_cell(row.cells[2], indicated, mono=True)
        _write_cell(row.cells[3], _format_decimal(r.computed_error), mono=True)
        _write_cell(row.cells[4], mpe_str, mono=True)
        _write_status_cell(row.cells[5], r.compliance_status)

    _apply_alternating_shading(table)
    _set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# Dispatch map: test_type -> builder function
TEST_BUILDERS = {
    TestType.WEIGHING_PERFORMANCE: _build_weighing_performance,
    TestType.ECCENTRICITY: _build_eccentricity,
    TestType.REPEATABILITY: _build_repeatability,
    TestType.DISCRIMINATION: _build_discrimination,
}


# ===================================================================
# Main entry point
# ===================================================================

def generate_docx(report) -> str:
    """
    Generate a NABL-style calibration certificate DOCX for *report*.

    Returns the absolute file path of the saved document.
    """
    context = _build_report_context(report)
    session = report.session
    instrument = session.instrument
    unit = instrument.unit
    results_qs = session.results.select_related('observation').order_by(
        'test_type', 'trial_number', 'test_point_load'
    )

    doc = Document()

    # ---------------------------------------------------------------
    # Global styles
    # ---------------------------------------------------------------
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.space_before = Pt(0)

    # Heading styles
    for level in (1, 2, 3):
        key = f'Heading {level}'
        if key in doc.styles:
            hs = doc.styles[key]
            hs.font.name = 'Calibri'
            hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # dark navy

    # ---------------------------------------------------------------
    # Page setup: A4, specified margins
    # ---------------------------------------------------------------
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)

    # ---------------------------------------------------------------
    # Header / Footer
    # ---------------------------------------------------------------
    _add_header_footer(doc, context)

    # ===============================================================
    # PAGE 1 --- Certificate Header + Instrument Details
    # ===============================================================

    # --- Optional lab logo (top-left) ---
    logo_data_uri = context.get('logo_data_uri', '')
    if logo_data_uri:
        try:
            import base64
            import io

            from docx.shared import Inches

            b64 = logo_data_uri.split(',', 1)[1]
            buf = io.BytesIO(base64.b64decode(b64))
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.add_run().add_picture(buf, height=Inches(0.6))
        except Exception:
            logger.exception("Logo decoding failed for DOCX; continuing without it")

    # --- Government header block (centred) ---
    _add_centered_text(doc, 'GOVERNMENT OF INDIA', size=12, bold=True,
                        uppercase=True, space_after=2, space_before=0)
    _add_centered_text(
        doc,
        context.get('jurisdiction',
                    'MINISTRY OF CONSUMER AFFAIRS, FOOD & PUBLIC DISTRIBUTION'),
        size=12, bold=True, uppercase=True, space_after=2)
    _add_centered_text(doc, 'DEPARTMENT OF CONSUMER AFFAIRS',
                        size=12, bold=True, uppercase=True, space_after=2)
    _add_centered_text(doc, 'LEGAL METROLOGY LABORATORY',
                        size=12, bold=True, space_after=2)
    _add_centered_text(doc, f"{context['lab_name']}, {context.get('lab_address', '')}",
                        size=11, bold=True, space_after=4)

    # Horizontal rule
    _add_horizontal_rule(doc)

    # Main title
    _add_centered_text(doc, 'TEST REPORT FOR NON-AUTOMATIC WEIGHING INSTRUMENT',
                        size=14, bold=True, space_before=6, space_after=4)
    _add_centered_text(
        doc,
        f"As per OIML R 76-1:2006  ·  {context['evaluation_type_label']}",
        size=10, space_after=4,
    )

    # Certificate number / ULR line
    p_cert = doc.add_paragraph()
    p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert.paragraph_format.space_after = Pt(2)
    run_cn = p_cert.add_run(f"Certificate No: {context['report_number']}")
    run_cn.font.size = Pt(10)
    run_cn.font.bold = True
    run_cn.font.name = 'Calibri'
    run_sep = p_cert.add_run('   /   ')
    run_sep.font.size = Pt(10)
    run_sep.font.name = 'Calibri'
    run_ulr = p_cert.add_run(f"ULR No: {context['accreditation_number']}")
    run_ulr.font.size = Pt(10)
    run_ulr.font.bold = True
    run_ulr.font.name = 'Calibri'

    # Issue date line
    formatted_date = _format_date_display(context['session_date'])
    _add_centered_text(doc, f'Certificate Issue Date: {formatted_date}',
                        size=10, space_after=8)

    # ---------------------------------------------------------------
    # Customer / Applicant Details
    # ---------------------------------------------------------------
    doc.add_heading('Customer / Applicant Details', level=2)

    cust_table = doc.add_table(rows=0, cols=2)
    cust_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _add_kv_row(cust_table, 'Customer / Applicant Name', context.get('lab_name', ''))
    _add_kv_row(cust_table, 'Address', context.get('lab_address', ''))
    _add_kv_row(cust_table, 'Contact Person', context.get('engineer_name', ''))
    _add_kv_row(cust_table, 'Test Request Date', formatted_date)
    _add_kv_row(cust_table, 'Date of Testing', formatted_date)
    _set_table_borders(cust_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------
    # Instrument Details
    # ---------------------------------------------------------------
    doc.add_heading('Instrument Under Test', level=2)

    inst_table = doc.add_table(rows=0, cols=2)
    inst_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _add_kv_row(inst_table, 'Instrument Detail', 'Non-Automatic Weighing Instrument')
    _add_kv_row(inst_table, 'Make / Manufacturer', instrument.manufacturer)
    _add_kv_row(inst_table, 'Model', instrument.model_name)
    _add_kv_row(inst_table, 'Serial Number', instrument.serial_number)
    _add_kv_row(inst_table, 'Accuracy Class',
                f'Class {instrument.accuracy_class}')
    _add_kv_row(inst_table, 'Max Capacity (Max)',
                f'{_format_decimal(instrument.max_capacity)} {unit}')
    _add_kv_row(inst_table, 'Min Capacity (Min)',
                f'{_format_decimal(instrument.min_capacity)} {unit}')
    _add_kv_row(inst_table, 'Verification Scale Interval (e)',
                f'{_format_decimal(instrument.verification_scale_interval_e)} {unit}')
    _add_kv_row(inst_table, 'Actual Scale Interval (d)',
                f'{_format_decimal(instrument.actual_scale_interval_d)} {unit}')
    _add_kv_row(inst_table, 'Number of Scale Intervals (n)',
                str(instrument.num_scale_intervals_n))
    tare_display = (instrument.tare_device_type or 'none').replace('_', ' ').title()
    _add_kv_row(inst_table, 'Tare Device', tare_display)
    if instrument.max_additive_tare:
        _add_kv_row(inst_table, 'Max Additive Tare',
                    f'{_format_decimal(instrument.max_additive_tare)} {unit}')
    if instrument.max_safe_load:
        _add_kv_row(inst_table, 'Max Safe Load (Lim)',
                    f'{_format_decimal(instrument.max_safe_load)} {unit}')
    if instrument.is_multi_interval:
        _add_kv_row(inst_table, 'Multi-Interval', 'Yes')
    _add_kv_row(inst_table, 'Instrument Location', context.get('lab_address', ''))
    _add_kv_row(inst_table, 'Equipment ID', instrument.serial_number)
    _set_table_borders(inst_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------
    # Environmental Conditions
    # ---------------------------------------------------------------
    doc.add_heading('Environmental Conditions During Test', level=2)

    env_table = doc.add_table(rows=0, cols=2)
    env_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Temperature: show range if both start/end available
    temp_start = context.get('temperature_start')
    temp_end = context.get('temperature_end')
    if temp_start is not None and temp_end is not None:
        _add_kv_row(env_table, 'Temperature',
                    f'{temp_start} °C to {temp_end} °C')
    elif temp_start is not None:
        _add_kv_row(env_table, 'Temperature', f'{temp_start} °C')
    elif temp_end is not None:
        _add_kv_row(env_table, 'Temperature', f'{temp_end} °C')

    if context.get('humidity') is not None:
        _add_kv_row(env_table, 'Relative Humidity', f"{context['humidity']} % RH")
    if context.get('barometric_pressure') is not None:
        _add_kv_row(env_table, 'Barometric Pressure',
                    f"{context['barometric_pressure']} hPa")
    _set_table_borders(env_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------
    # Important Remarks (bordered box with numbered items)
    # ---------------------------------------------------------------
    doc.add_heading('Important Remarks', level=2)

    remarks_table = doc.add_table(rows=1, cols=1)
    remarks_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = remarks_table.rows[0].cells[0]
    cell.text = ''
    remarks_list = context.get('remarks') or IMPORTANT_REMARKS
    for idx, remark in enumerate(remarks_list, start=1):
        p = cell.add_paragraph() if idx > 1 else cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(f'{idx}. {remark}')
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        run.font.color.rgb = GRAY_COLOR
    _set_cell_shading(cell, 'FFFFF0')
    _set_table_borders(remarks_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ---------------------------------------------------------------
    # Approval signature (right-aligned on page 1)
    # ---------------------------------------------------------------
    p_approval = doc.add_paragraph()
    p_approval.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_approval.paragraph_format.space_before = Pt(20)
    run_a1 = p_approval.add_run('Approved by:\n')
    run_a1.font.size = Pt(9)
    run_a1.font.name = 'Calibri'
    run_a1.font.bold = True
    approved_name = context.get('approved_by_name') or '________________________'
    run_a2 = p_approval.add_run(f'{approved_name}\n')
    run_a2.font.size = Pt(10)
    run_a2.font.name = 'Calibri'
    approved_date = _format_date_display(context.get('approved_at', ''))
    run_a3 = p_approval.add_run(f'Date: {approved_date}')
    run_a3.font.size = Pt(9)
    run_a3.font.name = 'Calibri'

    # ===============================================================
    # PAGE 2+ --- Test Results
    # ===============================================================
    doc.add_page_break()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    run_t = p_title.add_run('TEST RESULTS')
    run_t.font.size = Pt(14)
    run_t.font.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ref.paragraph_format.space_after = Pt(12)
    run_ref = p_ref.add_run(
        f"Report No: {context['report_number']}  |  "
        f"Instrument: {instrument.manufacturer} {instrument.model_name}  |  "
        f"S/N: {instrument.serial_number}"
    )
    run_ref.font.size = Pt(9)
    run_ref.font.name = 'Calibri'
    run_ref.font.color.rgb = GRAY_COLOR

    _add_horizontal_rule(doc)

    # Collect results by test type
    results_by_type: dict[str, list] = defaultdict(list)
    for r in results_qs:
        results_by_type[r.test_type].append(r)

    section_num = 1
    for test_type_value, label in TestType.choices:
        title = TEST_TYPE_TITLES.get(test_type_value, label)
        results = results_by_type.get(test_type_value, [])

        # Section heading
        h = doc.add_heading(f'{section_num}. {title}', level=3)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        section_num += 1

        if not results:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('Not Applicable — No observations recorded for this test.')
            run.font.color.rgb = GRAY_COLOR
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            continue

        # Dispatch to the appropriate builder
        builder = TEST_BUILDERS.get(test_type_value, _build_generic_test)
        builder(doc, results, unit, title)

    # ===============================================================
    # LAST PAGE --- Conclusion + Signatures
    # ===============================================================
    doc.add_page_break()

    # ---------------------------------------------------------------
    # Overall Test Summary Table
    # ---------------------------------------------------------------
    doc.add_heading('Overall Test Summary', level=2)

    # Build status map: worst result per test type
    status_map: dict[str, str] = {}
    for r in results_qs:
        tt = r.test_type
        cs = r.compliance_status
        if tt not in status_map:
            status_map[tt] = cs
        elif cs == ComplianceStatus.FAIL:
            status_map[tt] = ComplianceStatus.FAIL

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = summary_table.rows[0]
    _write_cell(hdr.cells[0], 'Sr No', size=9, bold=True)
    _write_cell(hdr.cells[1], 'Test Performed', size=9, bold=True)
    _write_cell(hdr.cells[2], 'Result', size=9, bold=True)
    _style_header_row(hdr)

    sr_no = 1
    for test_type_value, label in TestType.choices:
        s = status_map.get(test_type_value)
        row = summary_table.add_row()
        _write_cell(row.cells[0], str(sr_no), mono=True)
        _write_cell(row.cells[1], TEST_TYPE_TITLES.get(test_type_value, label),
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if s is None:
            _write_cell(row.cells[2], 'N/A', italic=True, color=GRAY_COLOR)
        else:
            _write_status_cell(row.cells[2], s)
        sr_no += 1

    _apply_alternating_shading(summary_table)
    _set_table_borders(summary_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ---------------------------------------------------------------
    # Conclusion Narrative
    # ---------------------------------------------------------------
    doc.add_heading('Conclusion', level=2)

    verdict = context['overall_verdict']
    if verdict == ComplianceStatus.PASS:
        conform_text = 'CONFORMS'
        verdict_color = PASS_COLOR
    else:
        conform_text = 'DOES NOT CONFORM'
        verdict_color = FAIL_COLOR

    p_conclusion = doc.add_paragraph()
    p_conclusion.paragraph_format.space_after = Pt(4)
    run_c1 = p_conclusion.add_run(
        'Based on the tests conducted as per OIML R 76-1:2006 '
        '(Non-automatic weighing instruments — Part 1: Metrological and technical '
        'requirements — Tests), the above instrument '
    )
    run_c1.font.size = Pt(10)
    run_c1.font.name = 'Calibri'

    run_c2 = p_conclusion.add_run(conform_text)
    run_c2.font.size = Pt(12)
    run_c2.font.bold = True
    run_c2.font.name = 'Calibri'
    run_c2.font.color.rgb = verdict_color

    run_c3 = p_conclusion.add_run(
        ' to the requirements specified in OIML R 76-1:2006 for '
        f'Accuracy Class {instrument.accuracy_class} instruments.'
    )
    run_c3.font.size = Pt(10)
    run_c3.font.name = 'Calibri'

    # Bold overall verdict banner
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    p_verdict = doc.add_paragraph()
    p_verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_verdict.paragraph_format.space_before = Pt(4)
    p_verdict.paragraph_format.space_after = Pt(12)
    if verdict == ComplianceStatus.PASS:
        run_v = p_verdict.add_run('✔  CONFORMS to OIML R 76-1:2006')
        run_v.font.color.rgb = PASS_COLOR
    else:
        run_v = p_verdict.add_run('✘  DOES NOT CONFORM to OIML R 76-1:2006')
        run_v.font.color.rgb = FAIL_COLOR
    run_v.font.size = Pt(14)
    run_v.font.bold = True
    run_v.font.name = 'Calibri'

    # ---------------------------------------------------------------
    # Three-column Signature Block
    # ---------------------------------------------------------------
    doc.add_heading('Signatories', level=2)

    sig_data = [
        {
            'role': 'Tested By',
            'title': 'Test Engineer',
            'name': context.get('engineer_name') or '________________________',
            'date': _format_date_display(context.get('session_date', '')),
        },
        {
            'role': 'Checked By',
            'title': 'Lab Manager',
            'name': context.get('checked_by_name') or '________________________',
            'date': '________________________',
        },
        {
            'role': 'Approved By',
            'title': 'Head of Laboratory',
            'name': context.get('approved_by_name') or '________________________',
            'date': _format_date_display(context.get('approved_at', '')),
        },
    ]

    sig_table = doc.add_table(rows=5, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_width = USABLE_WIDTH_CM / 3

    for col_idx, sig in enumerate(sig_data):
        # Row 0: Role
        cell_role = sig_table.rows[0].cells[col_idx]
        _write_cell(cell_role, sig['role'], size=10, bold=True)
        _set_cell_width(cell_role, col_width)

        # Row 1: Title in parentheses
        cell_title = sig_table.rows[1].cells[col_idx]
        _write_cell(cell_title, f"({sig['title']})", size=8, italic=True,
                    color=GRAY_COLOR)
        _set_cell_width(cell_title, col_width)

        # Row 2: Signature line (empty space)
        cell_sig = sig_table.rows[2].cells[col_idx]
        _write_cell(cell_sig, '\n\n________________________', size=9)
        _set_cell_width(cell_sig, col_width)

        # Row 3: Name
        cell_name = sig_table.rows[3].cells[col_idx]
        _write_cell(cell_name, sig['name'], size=10, bold=True)
        _set_cell_width(cell_name, col_width)

        # Row 4: Date
        cell_date = sig_table.rows[4].cells[col_idx]
        _write_cell(cell_date, f"Date: {sig['date']}", size=9)
        _set_cell_width(cell_date, col_width)

    # Light borders only on the signature table
    _set_table_borders(sig_table, color='CCCCCC', size='2')

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ---------------------------------------------------------------
    # QR code for public certificate verification
    # ---------------------------------------------------------------
    verify_url = context.get('verify_url')
    if verify_url:
        try:
            import io

            import qrcode
            from docx.shared import Inches

            qr = qrcode.QRCode(border=1, box_size=6)
            qr.add_data(verify_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color='black', back_color='white')
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            buf.seek(0)

            p_qr = doc.add_paragraph()
            p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_qr.add_run().add_picture(buf, width=Inches(0.9))

            p_qr_cap = doc.add_paragraph()
            p_qr_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_qr_cap.add_run(
                f"Scan to verify the authenticity of this certificate\n{verify_url}"
            )
            run_cap.font.size = Pt(7)
            run_cap.font.color.rgb = GRAY_COLOR
            run_cap.font.name = 'Calibri'
        except Exception:
            logger.exception("QR code generation failed for DOCX; continuing without it")

    # ---------------------------------------------------------------
    # Document Control Footer
    # ---------------------------------------------------------------
    _add_horizontal_rule(doc)

    p_doc_ctrl = doc.add_paragraph()
    p_doc_ctrl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_doc_ctrl.paragraph_format.space_before = Pt(4)
    run_dc = p_doc_ctrl.add_run(
        f"Doc No: {context.get('doc_control_number', DOC_CONTROL_NO)}  |  "
        f"Issue No: {context.get('doc_issue_number', DOC_ISSUE_NO)}  |  "
        f"Issue Date: {DOC_ISSUE_DATE}  |  "
        f"Rev No: {context.get('doc_rev_number', DOC_REV_NO)}"
    )
    run_dc.font.size = Pt(8)
    run_dc.font.color.rgb = GRAY_COLOR
    run_dc.font.name = 'Calibri'

    # Software attribution
    p_sw = doc.add_paragraph()
    p_sw.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sw.paragraph_format.space_before = Pt(4)
    run_sw = p_sw.add_run(
        f"Generated by NAWI Test Report Generator "
        f"v{context.get('software_version', '1.0')}  |  "
        "Standard reference: OIML R 76-1:2006"
    )
    run_sw.font.size = Pt(7)
    run_sw.font.color.rgb = GRAY_COLOR
    run_sw.font.name = 'Calibri'

    # ===============================================================
    # Save document
    # ===============================================================
    storage_dir = Path(settings.REPORT_STORAGE_PATH)
    storage_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{report.report_number.replace('/', '_')}_v{report.version}.docx"
    filepath = storage_dir / filename
    doc.save(str(filepath))

    logger.info("DOCX generated: %s", filepath)
    return str(filepath)
