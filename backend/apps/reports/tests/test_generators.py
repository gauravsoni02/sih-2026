import shutil
import tempfile
from decimal import Decimal
from pathlib import Path
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.test import TestCase, override_settings

from apps.instruments.models import Instrument
from apps.laboratory.models import Laboratory
from apps.reports.models import Report
from apps.testing.models import TestObservation, TestResult, TestSession

User = get_user_model()


def _create_test_fixtures():
    lab = Laboratory.objects.create(
        name='Central Metrology Lab',
        address='123 Test St, New Delhi',
        accreditation_number='NABL-TC-12345',
        lab_code='CML',
    )
    engineer = User.objects.create_user(
        username='eng_test', password='testpass123',
        role='engineer', laboratory=lab,
        first_name='Raj', last_name='Kumar',
    )
    instrument = Instrument.objects.create(
        manufacturer='Mettler Toledo',
        model_name='ICS465',
        serial_number='B734521098',
        accuracy_class='III',
        max_capacity=Decimal('30000'),
        min_capacity=Decimal('400'),
        verification_scale_interval_e=Decimal('20'),
        actual_scale_interval_d=Decimal('20'),
        num_scale_intervals_n=1500,
        unit='g',
        tare_device_type='additive',
        max_additive_tare=Decimal('30000'),
        max_safe_load=Decimal('45000'),
    )
    session = TestSession.objects.create(
        instrument=instrument,
        laboratory=lab,
        engineer=engineer,
        session_date='2026-08-31',
        temperature_start=Decimal('23.2'),
        temperature_end=Decimal('23.5'),
        humidity=Decimal('48.5'),
        barometric_pressure=Decimal('1013.2'),
        evaluation_type='initial_verification',
        verification_type='initial',
        status='completed',
        overall_verdict='pass',
    )

    wp_loads = [
        (Decimal('400'), Decimal('400')),
        (Decimal('6000'), Decimal('6001')),
        (Decimal('12000'), Decimal('12002')),
        (Decimal('21000'), Decimal('21001')),
        (Decimal('30000'), Decimal('30002')),
    ]
    for i, (load, indicated) in enumerate(wp_loads, 1):
        obs = TestObservation.objects.create(
            session=session,
            test_type='weighing_performance',
            test_point_load=load,
            indicated_value=indicated,
            correction=Decimal('0'),
            trial_number=i,
            direction='increasing',
        )
        error = indicated - load
        mpe = Decimal('10') if load <= Decimal('10000') else Decimal('20')
        TestResult.objects.create(
            session=session,
            test_type='weighing_performance',
            observation=obs,
            test_point_load=load,
            computed_error=error,
            mpe_applicable=mpe,
            compliance_status='pass',
            trial_number=i,
        )

    ecc_load = Decimal('20000')
    positions = ['center', 'front_left', 'front_right', 'rear_left', 'rear_right']
    readings = [ecc_load, ecc_load + 1, ecc_load - 1, ecc_load + 2, ecc_load - 2]
    for i, (pos, ind) in enumerate(zip(positions, readings), 1):
        obs = TestObservation.objects.create(
            session=session,
            test_type='eccentricity',
            test_point_load=ecc_load,
            indicated_value=ind,
            correction=Decimal('0'),
            position=pos,
            trial_number=i,
        )
        error = ind - ecc_load
        TestResult.objects.create(
            session=session,
            test_type='eccentricity',
            observation=obs,
            test_point_load=ecc_load,
            computed_error=error,
            mpe_applicable=Decimal('20'),
            compliance_status='pass',
            position=pos,
            trial_number=i,
        )

    rep_load = Decimal('15000')
    rep_readings = [rep_load, rep_load + 1, rep_load, rep_load, rep_load + 1, rep_load]
    for i, ind in enumerate(rep_readings, 1):
        obs = TestObservation.objects.create(
            session=session,
            test_type='repeatability',
            test_point_load=rep_load,
            indicated_value=ind,
            correction=Decimal('0'),
            trial_number=i,
        )
        TestResult.objects.create(
            session=session,
            test_type='repeatability',
            observation=obs,
            test_point_load=rep_load,
            computed_error=ind - rep_load,
            mpe_applicable=Decimal('20'),
            compliance_status='pass',
            trial_number=i,
        )

    report = Report.objects.create(
        report_number=Report.generate_report_number('CML'),
        session=session,
        generated_by=engineer,
        overall_verdict='pass',
        version=1,
    )
    return lab, engineer, instrument, session, report


class TestReportContextBuilder(TestCase):
    def setUp(self) -> None:
        self.lab, self.engineer, self.instrument, self.session, self.report = _create_test_fixtures()

    def test_build_report_context(self) -> None:
        from apps.reports.generators import _build_report_context
        ctx = _build_report_context(self.report)

        self.assertEqual(ctx['report_number'], self.report.report_number)
        self.assertEqual(ctx['lab_name'], 'Central Metrology Lab')
        self.assertEqual(ctx['accreditation_number'], 'NABL-TC-12345')
        self.assertEqual(ctx['engineer_name'], 'Raj Kumar')
        self.assertEqual(ctx['overall_verdict'], 'pass')
        self.assertEqual(ctx['evaluation_type_label'], 'Initial Verification')
        self.assertEqual(ctx['verification_type_label'], 'Initial')
        self.assertIsNotNone(ctx['instrument'])
        self.assertEqual(float(ctx['temperature_start']), 23.2)

    def test_context_with_approved_report(self) -> None:
        from apps.reports.generators import _build_report_context
        from django.utils import timezone

        manager = User.objects.create_user(
            username='mgr_test', password='testpass123',
            role='lab_manager', laboratory=self.lab,
            first_name='Priya', last_name='Singh',
        )
        self.report.approved_by = manager
        self.report.approved_at = timezone.now()
        self.report.save(update_fields=['approved_by', 'approved_at', 'updated_at'])

        ctx = _build_report_context(self.report)
        self.assertEqual(ctx['approved_by_name'], 'Priya Singh')
        self.assertTrue(len(ctx['approved_at']) > 0)


class TestPDFGenerator(TestCase):
    def setUp(self) -> None:
        self.lab, self.engineer, self.instrument, self.session, self.report = _create_test_fixtures()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_pdf_test_sections_built(self) -> None:
        from apps.reports.generators.pdf import _build_test_sections

        results_qs = self.session.results.all()
        sections = _build_test_sections(results_qs)

        self.assertGreater(len(sections), 0)

        wp_section = next((s for s in sections if s['title'] == 'Weighing Performance Test'), None)
        self.assertIsNotNone(wp_section)
        self.assertEqual(len(wp_section['results']), 5)
        self.assertEqual(wp_section['test_type'], 'weighing_performance')

        ecc_section = next((s for s in sections if s['title'] == 'Eccentricity Test'), None)
        self.assertIsNotNone(ecc_section)
        self.assertEqual(len(ecc_section['results']), 5)
        self.assertEqual(ecc_section['test_type'], 'eccentricity')

        rep_section = next((s for s in sections if s['title'] == 'Repeatability Test'), None)
        self.assertIsNotNone(rep_section)
        self.assertEqual(len(rep_section['results']), 6)
        self.assertEqual(rep_section['test_type'], 'repeatability')

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_pdf_compliance_summary(self) -> None:
        from apps.reports.generators.pdf import _build_compliance_summary

        results_qs = self.session.results.all()
        summary = _build_compliance_summary(results_qs)

        self.assertGreater(len(summary), 0)
        for item in summary:
            self.assertIn('test_name', item)
            self.assertIn('status', item)
            self.assertIn('status_label', item)

        wp_item = next((s for s in summary if 'Weighing' in s['test_name']), None)
        self.assertIsNotNone(wp_item)
        self.assertEqual(wp_item['status'], 'pass')

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_pdf_generates_file(self) -> None:
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            from apps.reports.generators.pdf import generate_pdf
            path = generate_pdf(self.report)

        self.assertTrue(Path(path).exists())
        self.assertTrue(path.endswith('.pdf'))
        self.assertGreater(Path(path).stat().st_size, 0)

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_pdf_empty_test_sections(self) -> None:
        from apps.reports.generators.pdf import _build_test_sections

        empty_session = TestSession.objects.create(
            instrument=self.instrument,
            laboratory=self.lab,
            engineer=self.engineer,
            session_date='2026-09-01',
            status='completed',
            overall_verdict='pass',
        )
        sections = _build_test_sections(empty_session.results.all())
        self.assertGreater(len(sections), 0)
        for section in sections:
            self.assertEqual(len(section['results']), 0)


class TestDOCXGenerator(TestCase):
    def setUp(self) -> None:
        self.lab, self.engineer, self.instrument, self.session, self.report = _create_test_fixtures()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_docx_generates_file(self) -> None:
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            from apps.reports.generators.docx import generate_docx
            path = generate_docx(self.report)

        self.assertTrue(Path(path).exists())
        self.assertTrue(path.endswith('.docx'))
        self.assertGreater(Path(path).stat().st_size, 0)

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_docx_contains_report_number(self) -> None:
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            from apps.reports.generators.docx import generate_docx
            path = generate_docx(self.report)

        from docx import Document
        doc = Document(path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        self.assertIn('Test Report', full_text)
        self.assertIn('CONFORMS', full_text)

    @override_settings(REPORT_STORAGE_PATH=None)
    def test_docx_contains_instrument_info(self) -> None:
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            from apps.reports.generators.docx import generate_docx
            path = generate_docx(self.report)

        from docx import Document
        doc = Document(path)
        all_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + ' '

        self.assertIn('Mettler Toledo', all_text)
        self.assertIn('ICS465', all_text)
        self.assertIn('B734521098', all_text)


class TestReportGeneration(TestCase):
    def setUp(self) -> None:
        self.lab, self.engineer, self.instrument, self.session, self.report = _create_test_fixtures()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_report_number_auto_increments(self) -> None:
        num1 = Report.generate_report_number('CML')
        Report.objects.create(
            report_number=num1,
            session=self.session,
            generated_by=self.engineer,
            overall_verdict='pass',
            version=2,
        )
        num2 = Report.generate_report_number('CML')
        seq1 = int(num1.split('/')[-1])
        seq2 = int(num2.split('/')[-1])
        self.assertEqual(seq2, seq1 + 1)

    def test_report_number_format(self) -> None:
        num = Report.generate_report_number('CML')
        parts = num.split('/')
        self.assertEqual(parts[0], 'NAWI')
        self.assertEqual(parts[1], 'CML')
        self.assertEqual(len(parts[3]), 4)

    def test_approved_report_immutable(self) -> None:
        from django.core.exceptions import ValidationError

        self.report.status = 'approved'
        self.report.save(update_fields=['status', 'updated_at'])

        with self.assertRaises(ValidationError):
            self.report.overall_verdict = 'fail'
            self.report.save()

    def test_approved_report_rejects_non_whitelisted_update_fields(self) -> None:
        from django.core.exceptions import ValidationError

        self.report.status = 'approved'
        self.report.save(update_fields=['status', 'updated_at'])

        with self.assertRaises(ValidationError):
            self.report.overall_verdict = 'fail'
            self.report.save(update_fields=['overall_verdict', 'updated_at'])

    def test_approved_report_allows_whitelisted_update_fields(self) -> None:
        self.report.status = 'approved'
        self.report.save(update_fields=['status', 'updated_at'])

        self.report.pdf_path = '/tmp/regenerated.pdf'
        self.report.save(update_fields=['pdf_path', 'updated_at'])
        self.report.refresh_from_db()
        self.assertEqual(self.report.pdf_path, '/tmp/regenerated.pdf')

    def test_report_versioning(self) -> None:
        self.assertEqual(self.report.version, 1)
        report_v2 = Report.objects.create(
            report_number=Report.generate_report_number('CML'),
            session=self.session,
            generated_by=self.engineer,
            overall_verdict='pass',
            version=2,
        )
        self.assertEqual(report_v2.version, 2)
        self.assertEqual(Report.objects.filter(session=self.session).count(), 2)


class TestReportGenerateAPI(TestCase):
    def setUp(self) -> None:
        from django.core.cache import cache
        from rest_framework.test import APIClient
        cache.clear()  # reset throttle history between tests
        self.api_client = APIClient()
        self.lab, self.engineer, self.instrument, self.session, _ = _create_test_fixtures()
        Report.objects.filter(session=self.session).delete()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_generate_report_for_completed_session(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp.status_code, 201)
        self.assertIn('report_number', resp.data)
        self.assertEqual(resp.data['overall_verdict'], 'pass')
        self.assertEqual(resp.data['version'], 1)

    def test_generate_report_for_draft_session_fails(self) -> None:
        self.session.status = 'draft'
        self.session.save(update_fields=['status'])

        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('completed', resp.data['detail'])

    def test_generate_report_increments_version(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            resp1 = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
            resp2 = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp1.data['version'], 1)
        self.assertEqual(resp2.data['version'], 2)

    def test_generate_nonexistent_session_404(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.post('/api/reports/generate/99999/')
        self.assertEqual(resp.status_code, 404)

    def test_generate_report_for_failed_session_rejected(self) -> None:
        self.session.overall_verdict = 'fail'
        self.session.save(update_fields=['overall_verdict'])

        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('DOES NOT CONFORM', resp.data['detail'])

    def test_viewer_cannot_generate(self) -> None:
        viewer = User.objects.create_user(
            username='viewer_rpt', password='testpass123', role='viewer',
        )
        self.api_client.force_authenticate(user=viewer)
        resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp.status_code, 403)

    def test_download_pdf(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            gen_resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
            report_id = gen_resp.data['id']
            resp = self.api_client.get(f'/api/reports/{report_id}/download/pdf/')
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp['Content-Type'], 'application/pdf')

    def test_download_docx(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            gen_resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
            report_id = gen_resp.data['id']
            resp = self.api_client.get(f'/api/reports/{report_id}/download/docx/')
        self.assertEqual(resp.status_code, 200)
        self.assertIn('wordprocessingml', resp['Content-Type'])

    def test_download_no_file_404(self) -> None:
        report = Report.objects.create(
            report_number=Report.generate_report_number('CML'),
            session=self.session,
            generated_by=self.engineer,
            overall_verdict='pass',
        )
        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.get(f'/api/reports/{report.pk}/download/pdf/')
        self.assertEqual(resp.status_code, 404)

    def test_other_lab_engineer_cannot_generate(self) -> None:
        other_lab = Laboratory.objects.create(
            name='Other Lab', address='Elsewhere',
            accreditation_number='NABL-TC-99999', lab_code='OTH',
        )
        outsider = User.objects.create_user(
            username='eng_other', password='testpass123',
            role='engineer', laboratory=other_lab,
        )
        self.api_client.force_authenticate(user=outsider)
        resp = self.api_client.post(f'/api/reports/generate/{self.session.pk}/')
        self.assertEqual(resp.status_code, 403)


class TestReviewApprovalWorkflow(TestCase):
    def setUp(self) -> None:
        from django.core.cache import cache
        from rest_framework.test import APIClient
        cache.clear()
        self.api_client = APIClient()
        self.lab, self.engineer, self.instrument, self.session, self.report = _create_test_fixtures()
        self.manager = User.objects.create_user(
            username='mgr_wf', password='testpass123',
            role='lab_manager', laboratory=self.lab,
            first_name='Priya', last_name='Singh',
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_engineer_cannot_review(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.post(f'/api/reports/{self.report.pk}/review/')
        self.assertEqual(resp.status_code, 403)

    def test_manager_review_sets_checked_by(self) -> None:
        self.api_client.force_authenticate(user=self.manager)
        resp = self.api_client.post(f'/api/reports/{self.report.pk}/review/')
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.data['status'], 'reviewed')
        self.report.refresh_from_db()
        self.assertEqual(self.report.checked_by, self.manager)
        self.assertIsNotNone(self.report.checked_at)

    def test_review_twice_rejected(self) -> None:
        self.api_client.force_authenticate(user=self.manager)
        self.api_client.post(f'/api/reports/{self.report.pk}/review/')
        resp = self.api_client.post(f'/api/reports/{self.report.pk}/review/')
        self.assertEqual(resp.status_code, 400)

    def test_approve_requires_review(self) -> None:
        self.api_client.force_authenticate(user=self.manager)
        resp = self.api_client.post(f'/api/reports/{self.report.pk}/approve/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('reviewed', resp.data['detail'])

    def test_approve_after_review_regenerates_files(self) -> None:
        self.api_client.force_authenticate(user=self.manager)
        self.api_client.post(f'/api/reports/{self.report.pk}/review/')
        with self.settings(REPORT_STORAGE_PATH=self.temp_dir):
            resp = self.api_client.post(f'/api/reports/{self.report.pk}/approve/')
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.data['status'], 'approved')
        self.report.refresh_from_db()
        self.assertEqual(self.report.status, 'approved')
        self.assertEqual(self.report.approved_by, self.manager)
        # Approval re-generates the certificate so the signed file carries
        # the approver and checked-by names.
        self.assertTrue(self.report.pdf_path)
        self.assertTrue(Path(self.report.pdf_path).exists())
        self.assertTrue(self.report.docx_path)
        self.assertTrue(Path(self.report.docx_path).exists())

    def test_context_includes_checked_by(self) -> None:
        from django.utils import timezone

        from apps.reports.generators import _build_report_context

        self.report.checked_by = self.manager
        self.report.checked_at = timezone.now()
        self.report.save(update_fields=['checked_by', 'checked_at', 'updated_at'])
        ctx = _build_report_context(self.report)
        self.assertEqual(ctx['checked_by_name'], 'Priya Singh')

    def test_preview_includes_org_settings(self) -> None:
        self.api_client.force_authenticate(user=self.engineer)
        resp = self.api_client.get(f'/api/reports/{self.report.pk}/preview/')
        self.assertEqual(resp.status_code, 200)
        self.assertIn('org_settings', resp.data)
        org = resp.data['org_settings']
        for key in ('jurisdiction', 'doc_control_number', 'doc_issue_number',
                    'doc_rev_number', 'doc_issue_date', 'default_remarks',
                    'logo_data_uri'):
            self.assertIn(key, org)
        self.assertTrue(len(org['default_remarks']) > 0)
        self.assertIn('checked_by', resp.data['report'])
        self.assertIn('approved_at', resp.data['report'])
